"""
CourseDocParser — extracts raw inputs from a .docx study guide.

Extracts: title, course_id, learning_objectives, content_sample, images.
Images are stored with position, caption (from doc text only), and alt_text
(only if not an AI-generated placeholder). NO visual descriptions are inferred.
"""

import hashlib
import re
import zipfile
from pathlib import Path
from typing import Optional
from xml.etree import ElementTree as ET

from docx import Document


class CourseDocParser:
    """Extracts raw inputs A0 needs from one or more equal-priority .docx sources."""

    def __init__(
        self,
        docx_paths: Optional[list[str]] = None,
        to_outline_doc_path: Optional[str] = None,
        *,
        docx_path: Optional[str] = None,
        extra_docx_paths: Optional[list[str]] = None,
    ):
        paths: list[str] = [str(p) for p in (docx_paths or []) if p]
        if not paths and docx_path:
            paths = [str(docx_path)]
            paths.extend(str(p) for p in (extra_docx_paths or []) if p)
        if not paths:
            raise ValueError("At least one docx path is required")

        self._docx_paths = [Path(p) for p in paths]
        self._sources: list[tuple[Path, Document]] = [
            (path, Document(str(path))) for path in self._docx_paths
        ]
        self._to_outline_document = (
            Document(to_outline_doc_path) if to_outline_doc_path else None
        )
        # First document — convenience for single-doc helpers and TO fetch
        self._docx_path = self._docx_paths[0]
        self.doc = self._sources[0][1]
        self.paragraphs = self.doc.paragraphs

    @staticmethod
    def _normalize_whitespace(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").strip())

    # Heading that is clearly a subsection, not the document / course title.
    _SUBSECTION_HEADING_RE = re.compile(
        r"^\d+(\.\d+)*\s+("
        r"learning\s+objectives|objectives|overview|summary|introduction|"
        r"conclusion|appendix|references|bibliography|toc|table\s+of\s+contents"
        r")\b",
        re.IGNORECASE,
    )

    def _title_from_filename(self, path: Optional[Path] = None) -> str:
        stem = (path or self._docx_path).stem
        s = stem
        s = re.sub(r"(?i)_ACCEPTED|_FINAL|_DRAFT|_REV[A-Z0-9]*", "", s)
        s = re.sub(r"(?i)_SG_", " ", s)
        s = re.sub(r"_\d{6,8}(?![0-9])", "", s)
        s = re.sub(r"^\d+_", "", s)
        s = s.replace("_", " ")
        s = self._normalize_whitespace(s)
        if len(s) >= 3:
            return " ".join(w.capitalize() for w in s.split())
        return ""

    def _heading_might_be_course_title(self, text: str) -> bool:
        if len(text) < 8:
            return False
        if self._SUBSECTION_HEADING_RE.match(text):
            return False
        return True

    def _extract_title_from_doc(self, doc: Document, path: Path) -> str:
        """Best-effort title from one document."""
        cp = (doc.core_properties.title or "").strip()
        if cp and len(cp) > 2:
            return self._normalize_whitespace(cp)

        paragraphs = doc.paragraphs
        for p in paragraphs[:80]:
            if p.style.name == "Title" and p.text.strip():
                t = self._normalize_whitespace(p.text)
                if len(t) > 2:
                    return t

        for p in paragraphs[:40]:
            name = p.style.name
            if "Heading" not in name:
                continue
            if not p.text.strip():
                continue
            t = self._normalize_whitespace(p.text)
            if not self._heading_might_be_course_title(t):
                continue
            if name == "Heading 1" or name == "Heading 2" or len(t) >= 40:
                return t

        for p in paragraphs[:25]:
            t = self._normalize_whitespace(p.text)
            if len(t) < 25:
                continue
            if p.style.name in self._STOP_STYLES:
                continue
            if self._SUBSECTION_HEADING_RE.match(t):
                continue
            return t

        derived = self._title_from_filename(path)
        if derived:
            return derived
        return ""

    def extract_title(self) -> str:
        """
        Best-effort title across all source documents (first non-empty wins),
        then filename stem of the first file.
        """
        for path, doc in self._sources:
            title = self._extract_title_from_doc(doc, path)
            if title:
                return title
        return "Course"

    def extract_course_id(self) -> Optional[str]:
        for _path, doc in self._sources:
            for p in doc.paragraphs[:10]:
                m = re.search(r"Course\s*ID[:\s]*(\d+)", p.text, re.IGNORECASE)
                if m:
                    return m.group(1)
        return None

    # Phrases that signal the start of a learning-objectives block.
    _LO_TRIGGERS = (
        "learning objectives",
        "learning outcomes",
        "course objectives",
    )

    # Paragraph styles that mark the start of a NEW major section → stop capture.
    _STOP_STYLES = {"Title", "Heading 1", "Heading 2"}

    # Leading bullet / dingbat (optional spaces). Word often omits space after •.
    _BULLET_PREFIX_RE = re.compile(
        r"^[\s]*(?:"
        r"[\u2022\u2023\u25E6\u2043\u2219\u25AA\u25CF\u25CB\u2218\u00B7"
        r"\u25D8\u25BA\u27A4\u204C\u204D\-–—\*•▪▸►◦◇●]"
        r"[\s]*)+"
    )
    # Numbered list items: "1.", "1)", "(1)", "a.", "a)", "1 " (single digit + space)
    _NUMBERED_RE = re.compile(
        r"^(?:\(?[0-9]+[.)]\s+|\(?[a-zA-Z][.)]\s+|[0-9]{1,}\s+(?=\S))"
    )

    # First word of a line that usually starts a body paragraph, not an LO.
    _LO_SKIP_FIRST_WORD = frozenset(
        {
            "the", "this", "these", "those", "when", "if", "for", "there", "here",
            "it", "its", "as", "while", "however", "although", "because", "since",
            "during", "after", "before", "by", "in", "on", "at", "an", "and", "or",
            "but", "nor", "we", "you", "they", "he", "she", "our", "your", "their",
            "some", "many", "most", "all", "each", "every", "such", "one", "two",
            "three", "both", "another", "other", "any", "no", "not", "only", "also",
            "from", "into", "with", "without", "within", "about", "above",
        }
    )
    _LO_FIRST_TOKEN_RE = re.compile(
        r"^(?P<first>[A-Za-z]{2,}(?:'[sS])?)(?=[\s\.,;:\/\-]|$)"
    )

    def _strip_bullet_prefix(self, text: str) -> str:
        return self._BULLET_PREFIX_RE.sub("", text).strip()

    def _has_leading_bullet_marker(self, text: str) -> bool:
        return bool(self._BULLET_PREFIX_RE.match(text)) and bool(
            self._strip_bullet_prefix(text)
        )

    def _is_lo_intro_line(self, text: str, already_have_objectives: bool) -> bool:
        """Generic intro before bullet/prose objectives (e.g. ends with colon)."""
        if already_have_objectives or len(text) > 220:
            return False
        # Typical lead-in: "…you should be able to:" — no fixed phrase list.
        return text.rstrip().endswith(":")

    def _looks_like_prose_objective(self, text: str) -> bool:
        """Non-list line that still reads like an outcome (not a hardcoded verb list)."""
        t = text.strip()
        if len(t) < 15:
            return False
        if t.rstrip().endswith(";"):
            return True
        m = self._LO_FIRST_TOKEN_RE.match(t)
        if not m:
            return False
        first = m.group("first").lower()
        if first in self._LO_SKIP_FIRST_WORD:
            return False
        return len(t) >= 25

    def extract_learning_objectives(self, paragraphs=None) -> list[str]:
        """
        Robustly extracts the learning-objectives block from the study-guide doc.

        Strategy
        --------
        1. Scan ALL paragraphs for a trigger phrase (no depth limit).
        2. Once triggered, collect paragraphs that look like list items:
             • "List *" paragraph styles, OR
             • Text that starts with a bullet/dingbat (with or without space after), OR
             • Numbered list pattern, OR
             • Prose lines: end with ``;`` OR start with a non-body first word and are long enough.
        3. Short intro lines ending in ``:`` before the first objective are skipped
           (no fixed phrase list).
        4. Allow up to MAX_SKIP consecutive non-matching, non-empty paragraphs
           before stopping.
        5. Stop immediately if a new major heading (Heading 1/2, Title) is hit
           AFTER at least one objective has been collected.
        """
        if paragraphs is None:
            paragraphs = self.paragraphs

        MAX_SKIP = 3          # non-list paragraphs tolerated before giving up
        objectives: list[str] = []
        capture = False
        skip_count = 0

        for p in paragraphs:
            text  = p.text.strip()
            style = p.style.name
            low   = text.lower()

            # ── Check for trigger phrase (re-arm even mid-doc) ────────────────
            if any(phrase in low for phrase in self._LO_TRIGGERS):
                capture    = True
                skip_count = 0
                continue

            if not capture:
                continue

            # ── Stop on a new major heading (only once we have something) ─────
            if objectives and style in self._STOP_STYLES:
                break

            # ── Empty line → don't count against skip budget ──────────────────
            if not text:
                continue

            # ── Intro line before bullets/prose (e.g. "… you should be able to:")
            if self._is_lo_intro_line(text, bool(objectives)):
                continue

            # ── Decide whether this paragraph is a list item ──────────────────
            is_list_style = style.startswith("List") or style in (
                "List Paragraph",
                "List Number",
                "List Number 2",
            )
            is_bullet_text   = self._has_leading_bullet_marker(text)
            is_numbered_text = bool(self._NUMBERED_RE.match(text))

            if is_list_style or is_bullet_text or is_numbered_text:
                # Strip number / bullet prefix to get the objective text
                if is_numbered_text:
                    clean = self._NUMBERED_RE.sub("", text).strip()
                elif is_bullet_text:
                    clean = self._strip_bullet_prefix(text)
                else:
                    clean = text
                objectives.append(clean.rstrip(";").rstrip("."))
                skip_count = 0
            elif self._looks_like_prose_objective(text):
                objectives.append(text.rstrip(";").rstrip("."))
                skip_count = 0
            else:
                # Non-list paragraph (bridge sentence, sub-header, etc.)
                skip_count += 1
                if skip_count > MAX_SKIP:
                    # Consistent non-list content → we've left the LO block
                    break

        return objectives

    def extract_merged_learning_objectives(self) -> list[str]:
        """Extract and deduplicate learning objectives across all source documents."""
        all_objectives: list[str] = []
        seen: set[str] = set()

        for _path, doc in self._sources:
            for obj in self.extract_learning_objectives(paragraphs=doc.paragraphs):
                key = obj.lower()
                if key not in seen:
                    all_objectives.append(obj)
                    seen.add(key)

        return all_objectives

    def extract_merged_full_content(self, max_words: int = 8000) -> str:
        """Extract and merge body text from all source documents up to max_words."""
        parts: list[str] = []
        total_words = 0

        for path, doc in self._sources:
            if total_words >= max_words:
                break
            parts.append(f"\n--- Document: {path.name} ---\n")
            for p in doc.paragraphs:
                if total_words >= max_words:
                    parts.append("[…content truncated at word limit…]")
                    break
                text = p.text.strip()
                if not text:
                    continue
                words = text.split()
                if total_words + len(words) >= max_words:
                    remaining = max_words - total_words
                    parts.append(" ".join(words[:remaining]))
                    parts.append("[…content truncated at word limit…]")
                    total_words = max_words
                    break
                parts.append(text)
                total_words += len(words)

        return "\n".join(parts)

    def count_paragraphs(self) -> int:
        """Return total paragraph count across all source documents."""
        return sum(len(doc.paragraphs) for _path, doc in self._sources)

    def fetch_paragraphs_by_range(
        self,
        para_idx_start: Optional[int],
        para_idx_end: Optional[int],
        skip_empty: bool = True,
    ) -> list[str]:
        """Return the raw text of paragraphs para_idx_start … para_idx_end (inclusive).

        Uses the actual doc.paragraphs indices, matching the [P<N>] values stored
        in each TO section's para_idx_start / para_idx_end fields.

        Args:
            para_idx_start: First paragraph index (from TO section, may be None).
            para_idx_end:   Last paragraph index inclusive (from TO section, may be None).
            skip_empty:     If True, omit blank paragraphs from the result.

        Returns:
            List of paragraph text strings for that section's source range.
            Returns [] when either index is None or out of range.
        """
        if para_idx_start is None or para_idx_end is None:
            return []

        total = len(self.doc.paragraphs)
        start = max(0, para_idx_start)
        end = min(total - 1, para_idx_end)

        if start > end:
            return []

        texts: list[str] = []
        for p in self.doc.paragraphs[start : end + 1]:
            text = p.text.strip()
            if skip_empty and not text:
                continue
            texts.append(text)
        return texts

    def fetch_section_content(
        self,
        para_idx_start: Optional[int],
        para_idx_end: Optional[int],
        max_words: int = 5000,
    ) -> str:
        """Return joined paragraph text for a TO section's source range.

        Convenience wrapper around fetch_paragraphs_by_range() that joins
        paragraphs into a single string and enforces a word-count cap.

        Args:
            para_idx_start: First paragraph index (from TO section).
            para_idx_end:   Last paragraph index inclusive.
            max_words:      Cap on total words returned (prevents oversized LLM prompts).

        Returns:
            Single string of the section's raw source content.
        """
        paragraphs = self.fetch_paragraphs_by_range(para_idx_start, para_idx_end)
        if not paragraphs:
            return ""

        collected: list[str] = []
        total_words = 0
        for text in paragraphs:
            words = text.split()
            if total_words + len(words) > max_words:
                remaining = max_words - total_words
                collected.append(" ".join(words[:remaining]))
                collected.append("[…content truncated at word limit…]")
                break
            collected.append(text)
            total_words += len(words)

        return "\n".join(collected)

    def get_section_heading_map(self) -> list[tuple[int, str, int] | tuple[int, str, int, str]]:
        """Return headings from all source documents for TO section mapping.

        Each entry is (para_idx, heading_text, heading_level) or, when multiple
        files are loaded, (para_idx, heading_text, heading_level, source_filename).
        """
        multi = len(self._sources) > 1
        result: list[tuple[int, str, int] | tuple[int, str, int, str]] = []
        for path, doc in self._sources:
            for idx, p in enumerate(doc.paragraphs):
                if "Heading" not in p.style.name or not p.text.strip():
                    continue
                try:
                    level = int(p.style.name[-1]) if p.style.name[-1].isdigit() else 0
                except (IndexError, ValueError):
                    level = 0
                text = p.text.strip()
                if multi:
                    result.append((idx, text, level, path.name))
                else:
                    result.append((idx, text, level))
        return result

    def extract_heading_tree(
        self, paragraphs=None, source_label: str = "primary"
    ) -> list[dict]:
        """Return a structured heading tree for one document.

        Each entry: {level, text, para_idx, source}.

        Level 1 headings are top-level topics; level 2+ are sub-topics.
        Also detects numbered headings like "7.0 Topic" even when paragraph
        style is "Normal", since some DOCXs don't use Heading styles.
        """
        if paragraphs is None:
            paragraphs = self.doc.paragraphs

        _NUMBERED_HEADING = re.compile(
            r"^(\d+(?:\.\d+)*)[\s.\-:]\s*(.+)$"
        )
        result: list[dict] = []

        for idx, p in enumerate(paragraphs):
            text = p.text.strip()
            if not text:
                continue
            style = p.style.name

            # Named heading style
            if "Heading" in style:
                try:
                    level = int(style[-1]) if style[-1].isdigit() else 1
                except (IndexError, ValueError):
                    level = 1
                result.append({"level": level, "text": text, "para_idx": idx, "source": source_label})
                continue

            # Numbered section heading (e.g. "3.0 Overview", "3.1.2 Sub-topic")
            m = _NUMBERED_HEADING.match(text)
            if m:
                dots = m.group(1).count(".")
                level = dots + 1  # "3" → level 1, "3.1" → level 2, etc.
                result.append({"level": level, "text": text, "para_idx": idx, "source": source_label})

        return result

    def extract_merged_heading_tree(self) -> list[dict]:
        """Extract and merge heading trees from all source documents (deduplicated by text)."""
        tree: list[dict] = []
        seen: set[str] = set()

        for path, doc in self._sources:
            label = path.name
            for h in self.extract_heading_tree(paragraphs=doc.paragraphs, source_label=label):
                key = h["text"].lower()
                if key not in seen:
                    tree.append(h)
                    seen.add(key)

        return tree

    def extract_indexed_content(self, max_words: int = 8000) -> str:
        """Return all source documents with [P<N>] paragraph index markers per file.

        Each document block starts with ``--- Document: <filename> ---``.
        [P<N>] is the paragraph index within that document's doc.paragraphs list.
        """
        lines: list[str] = []
        total_words = 0

        for path, doc in self._sources:
            if total_words >= max_words:
                break
            lines.append(f"\n--- Document: {path.name} ---")
            for idx, p in enumerate(doc.paragraphs):
                if total_words >= max_words:
                    lines.append("[…content truncated at word limit…]")
                    break
                text = p.text.strip()
                if not text:
                    continue
                words = text.split()
                if total_words + len(words) > max_words:
                    remaining = max_words - total_words
                    lines.append(f"[P{idx}] " + " ".join(words[:remaining]))
                    lines.append("[…content truncated at word limit…]")
                    total_words = max_words
                    break
                lines.append(f"[P{idx}] {text}")
                total_words += len(words)

        return "\n".join(lines)

    def extract_content_sample(self, max_chars: int = 3000) -> str:
        """High-level sample from all sources: headings + first paragraph per section."""
        parts: list[str] = []
        total = 0
        for path, doc in self._sources:
            if total >= max_chars:
                break
            parts.append(f"\n--- Document: {path.name} ---")
            total += len(parts[-1])
            prev_was_heading = False
            for p in doc.paragraphs:
                if "Heading" in p.style.name and p.text.strip():
                    parts.append(f"\n[{p.style.name}] {p.text.strip()}")
                    total += len(parts[-1])
                    prev_was_heading = True
                elif prev_was_heading and p.text.strip():
                    snippet = p.text.strip()
                    parts.append(snippet)
                    total += len(snippet)
                    prev_was_heading = False
                if total >= max_chars:
                    break
        return "\n".join(parts)

    def extract_full_content(self, max_words: int = 8000) -> str:
        """Extract raw body text from the entire document up to max_words.

        Includes all paragraph text in document order so the LLM gets real
        course content rather than a structural outline of headings only.
        """
        collected: list[str] = []
        word_count = 0
        for p in self.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            words = text.split()
            if word_count + len(words) >= max_words:
                remaining = max_words - word_count
                collected.append(" ".join(words[:remaining]))
                collected.append("[…content truncated at word limit…]")
                break
            collected.append(text)
            word_count += len(words)
        return "\n".join(collected)

    def extract_all_images(self, images_dir: Path) -> list[dict]:
        """Extract embedded images from every source document."""
        images_dir.mkdir(parents=True, exist_ok=True)
        all_images: list[dict] = []
        seq_offset = 0
        for path, doc in self._sources:
            self.doc = doc
            self.paragraphs = doc.paragraphs
            batch = self._extract_images_from_doc(
                str(path), images_dir, start_seq=seq_offset
            )
            for img in batch:
                img["source_document"] = path.name
                all_images.append(img)
            seq_offset += len(batch)
        return all_images

    def extract_images(self, docx_path: str, images_dir: Path) -> list[dict]:
        """Extract images from a single docx (backward-compatible alias)."""
        return self._extract_images_from_doc(docx_path, images_dir)

    def _extract_images_from_doc(
        self, docx_path: str, images_dir: Path, *, start_seq: int = 0
    ) -> list[dict]:
        """
        Extract all embedded images from one docx.

        Stores:
          - binary file to images_dir
          - position (para_idx)
          - caption: ONLY from explicit surrounding text (prev/next paragraph)
          - alt_text: ONLY if not AI-generated placeholder
          - heading context at time of appearance

        NO visual descriptions are inferred by LLM or code.
        """
        images_dir.mkdir(parents=True, exist_ok=True)

        # Build rId -> media filename map from relationships XML
        rid_to_media: dict[str, str] = {}
        with zipfile.ZipFile(docx_path) as z:
            try:
                rels_xml = z.read("word/_rels/document.xml.rels")
            except KeyError:
                return []
            root = ET.fromstring(rels_xml)
            ns = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}
            for rel in root.findall("r:Relationship", ns):
                if "image" in rel.get("Type", "").lower():
                    rid_to_media[rel.get("Id")] = rel.get("Target", "").replace("media/", "")

        # Namespace shortcuts
        WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
        R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

        images: list[dict] = []
        img_seq = start_seq
        heading_ctx = ("", 0)

        for para_idx, p in enumerate(self.paragraphs):
            style = p.style.name
            text  = p.text.strip()

            # Track heading context
            if "Heading" in style and text:
                level = int(style[-1]) if style[-1].isdigit() else 0
                heading_ctx = (text, level)

            pxml = p._p

            # Find all drawing containers (inline + anchored)
            for container_tag in (f"{{{WP}}}inline", f"{{{WP}}}anchor"):
                for container in pxml.iter(container_tag):

                    # docPr — name and alt text
                    doc_pr = container.find(f"{{{WP}}}docPr")
                    raw_descr = (doc_pr.get("descr", "") if doc_pr is not None else "")

                    # Strip AI-generated disclaimer — do NOT use as description
                    ai_disclaimer = "AI-generated content may be incorrect"
                    alt_text = (
                        raw_descr.replace(ai_disclaimer, "").strip()
                        if ai_disclaimer not in raw_descr
                        else ""
                    )

                    # Size in cm (EMUs: 1 cm = 360000 EMU)
                    extent = container.find(f"{{{WP}}}extent")
                    width_cm  = round(int(extent.get("cx", 0)) / 360000, 1) if extent is not None else None
                    height_cm = round(int(extent.get("cy", 0)) / 360000, 1) if extent is not None else None

                    # r:embed -> rId
                    blip = container.find(f".//{{{A}}}blip")
                    if blip is None:
                        continue
                    r_embed = blip.get(f"{{{R}}}embed", "")
                    media_filename = rid_to_media.get(r_embed, "")
                    if not media_filename:
                        continue

                    # Read binary and save
                    try:
                        with zipfile.ZipFile(docx_path) as z:
                            img_bytes = z.read(f"word/media/{media_filename}")
                    except KeyError:
                        continue

                    ext       = Path(media_filename).suffix
                    img_seq  += 1
                    img_id    = f"img_{img_seq:03d}"
                    save_name = f"{img_id}_{media_filename}"
                    save_path = images_dir / save_name
                    save_path.write_bytes(img_bytes)

                    # Caption detection — only from explicit preceding/following text
                    prev_text = (self.paragraphs[para_idx - 1].text.strip()
                                 if para_idx > 0 else "")
                    next_text = (self.paragraphs[para_idx + 1].text.strip()
                                 if para_idx < len(self.paragraphs) - 1 else "")

                    intro_triggers = (
                        "following", "below", "illustrated", "shown", "depicts",
                        "as seen", "figure", "chart", "graph", "table", "map",
                        "image", "diagram"
                    )
                    caption = ""
                    if any(t in prev_text.lower() for t in intro_triggers):
                        caption = prev_text
                    elif next_text and len(next_text) <= 120 and not next_text.endswith("."):
                        caption = next_text

                    images.append({
                        "id": img_id,
                        "r_embed": r_embed,
                        "media_filename": media_filename,
                        "saved_path": str(save_path),
                        "para_idx": para_idx,
                        "size_cm": {"width": width_cm, "height": height_cm},
                        "size_bytes": len(img_bytes),
                        "sha256": hashlib.sha256(img_bytes).hexdigest()[:16],
                        "caption": caption,
                        "has_caption": bool(caption),
                        "alt_text": alt_text,
                        "heading_context": heading_ctx[0],
                        "heading_level": heading_ctx[1],
                    })

        return images

    def count_total_doc_words(self) -> int:
        """Return the total word count across all source documents."""
        total = 0
        for _path, doc in self._sources:
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    total += len(text.split())
        return total

    def extract_to_outline_text(self) -> str:
        """
        Extracts TO document into clean, structure-agnostic text.

        Works for:
        - paragraphs
        - tables (if present)
        - mixed / broken formats

        Goal: maximize LLM understanding, NOT preserve formatting
        """
        if self._to_outline_document is None:
            return ""

        chunks = []

        def clean(text: str) -> str:
            return re.sub(r"\s+", " ", text).strip()

        # ---- 1. Extract ALL paragraphs ----
        for p in self._to_outline_document.paragraphs:
            text = clean(p.text)
            if not text:
                continue

            style = p.style.name.lower()

            # Detect headings flexibly
            if (
                "heading" in style
                or text.isupper()
                or any(
                    k in text.lower()
                    for k in [
                        "course",
                        "lesson",
                        "module",
                        "section",
                        "objective",
                        "outline",
                        "conclusion",
                    ]
                )
            ):
                chunks.append(f"\n[SECTION] {text}")
            else:
                chunks.append(text)

        # ---- 2. Extract tables IF present (optional support) ----
        for table in self._to_outline_document.tables:
            for row in table.rows:
                row_data = []

                for cell in row.cells:
                    cell_text = clean(cell.text)
                    if cell_text:
                        row_data.append(cell_text)

                if row_data:
                    # Convert row → sentence instead of rigid format
                    sentence = " | ".join(row_data)
                    chunks.append(f"[ROW] {sentence}")

        # ---- 3. Final normalization ----
        text_output = "\n".join(chunks)

        # Remove excessive newlines
        text_output = re.sub(r"\n{3,}", "\n\n", text_output)

        return text_output.strip()
