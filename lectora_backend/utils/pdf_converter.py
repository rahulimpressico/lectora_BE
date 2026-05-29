"""
PDF → DOCX conversion utility.

This module intentionally uses only ``pypdf`` + ``python-docx``.

It performs a text-first conversion that:
- injects PDF bookmark entries as ``TOC N`` paragraphs when available
- detects numbered / ALL-CAPS headings heuristically
- preserves bullet and numbered list items when possible
- falls back to normal body paragraphs for all other extracted lines
"""
from __future__ import annotations

import logging
import re
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)

# ── List / heading detection regexes ─────────────────────────────────────────
# Common bullet characters (Unicode + ASCII)
_BULLET_RE = re.compile(
    r"^[•‣◦⁃∙▪▫●○"
    r"✔✓▸▹►▻➔➢•●◦○▪▸▶\-–—]\s+"
)
_NUMBER_RE = re.compile(r"^\s*\d+[.)]\s+")
_NUMBERED_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)[.\s]\s*(.+)$")
_ALL_CAPS_RE = re.compile(r"^[A-Z][A-Z0-9 \-/&,.]{3,79}$")


# ── Public API ────────────────────────────────────────────────────────────────

def convert_pdf_to_docx_bytes(pdf_bytes: bytes, stem: str = "document") -> bytes:
    """Convert raw PDF bytes to DOCX bytes using only pypdf + python-docx."""
    safe_stem = re.sub(r"[^\w\-]", "_", stem)[:80] or "document"

    with tempfile.TemporaryDirectory(prefix="lectora_pdf_docx_") as tmp:
        tmp_dir  = Path(tmp)
        pdf_path  = tmp_dir / f"{safe_stem}.pdf"
        docx_path = tmp_dir / f"{safe_stem}.docx"
        pdf_path.write_bytes(pdf_bytes)

        return _pypdf_conversion(pdf_path, docx_path, safe_stem)


def _add_toc_section(toc: list[tuple[int, str, int | None]], docx_doc) -> None:
    """Add a TOC heading + ``TOC N``-styled entries (document is empty so far)."""
    from docx.shared import Pt

    docx_doc.add_heading("Table of Contents", level=1)

    for level, title, _page in toc:
        style_name = f"TOC {min(max(level, 1), 9)}"
        try:
            p = docx_doc.add_paragraph(title, style=style_name)
        except (KeyError, Exception):
            p = docx_doc.add_paragraph(title)
            p.paragraph_format.left_indent = Pt(12 * (level - 1))

    docx_doc.add_paragraph("")  # visual separator after TOC


def _list_type(text: str) -> "str | None":
    """Return ``'bullet'``, ``'number'``, or ``None``."""
    if _BULLET_RE.match(text):
        return "bullet"
    if _NUMBER_RE.match(text):
        return "number"
    return None


def _extract_outline_entries(reader) -> list[tuple[int, str, int | None]]:
    """Return PDF bookmark entries as ``(level, title, page_number)`` tuples."""
    outline_items = getattr(reader, "outline", None)
    if not outline_items:
        return []

    entries: list[tuple[int, str, int | None]] = []

    def walk(items, level: int = 1) -> None:
        for item in items:
            if isinstance(item, list):
                walk(item, level + 1)
                continue

            title = str(getattr(item, "title", "") or item).strip()
            if not title:
                continue

            page_number = None
            try:
                page_number = reader.get_destination_page_number(item) + 1
            except Exception:
                page_number = None

            entries.append((level, title, page_number))

    try:
        walk(outline_items, 1)
    except Exception as exc:
        logger.debug("[pdf_converter] Could not read PDF outline: %s", exc)
        return []

    return entries


def _pypdf_conversion(pdf_path: Path, docx_path: Path, stem: str) -> bytes:
    """Extract text via pypdf and build a structured DOCX with heading heuristics.

    Heading detection:
    - Numbered headings (e.g. ``"3.1  Sub-section Title"``)
    - ALL-CAPS short lines (≤ 80 chars — common section-header pattern)

    Everything else becomes a body paragraph.
    """
    try:
        from pypdf import PdfReader           # type: ignore[import-untyped]
        from docx import Document             # type: ignore[import-untyped]
    except ImportError as exc:
        raise RuntimeError(
            f"pypdf/python-docx are required to convert PDF to DOCX: {exc}"
        ) from exc

    reader = PdfReader(str(pdf_path))
    doc = Document()
    doc.core_properties.title = stem.replace("_", " ").title()
    toc = _extract_outline_entries(reader)
    if toc:
        _add_toc_section(toc, doc)

    prev_blank = False

    for page in reader.pages:
        raw = (page.extract_text() or "").strip()
        if not raw:
            continue

        for line in raw.splitlines():
            text = line.strip()
            if not text:
                if not prev_blank:
                    doc.add_paragraph("")
                prev_blank = True
                continue
            prev_blank = False

            # Numbered heading — derive level from dot count
            m = _NUMBERED_HEADING_RE.match(text)
            if m:
                dots = m.group(1).count(".")
                doc.add_heading(text, level=min(dots + 1, 6))
                continue

            # ALL-CAPS short line → top-level heading
            if len(text) <= 80 and _ALL_CAPS_RE.match(text):
                doc.add_heading(text, level=1)
                continue

            list_type = _list_type(text)
            if list_type == "bullet":
                doc.add_paragraph(_BULLET_RE.sub("", text).strip(), style="List Bullet")
                continue
            if list_type == "number":
                doc.add_paragraph(_NUMBER_RE.sub("", text).strip(), style="List Number")
                continue

            doc.add_paragraph(text)

        doc.add_paragraph("")  # lightweight page separator

    doc.save(str(docx_path))
    result = docx_path.read_bytes()
    logger.info(
        "[pdf_converter] pypdf conversion OK: %s → %d bytes", stem, len(result)
    )
    return result
